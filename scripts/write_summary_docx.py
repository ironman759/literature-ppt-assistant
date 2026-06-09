#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


LATIN_FONT = "Times New Roman"
CJK_BODY_FONT = "Songti SC"
CJK_HEADING_FONT = "Heiti SC"


def has_cjk(text: str) -> bool:
    return any("\u3400" <= ch <= "\u9fff" for ch in text)


def set_east_asia_font(run, east_asia: str, latin: str = LATIN_FONT) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def set_style_font(style, east_asia: str, latin: str = LATIN_FONT) -> None:
    style.font.name = latin
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def add_page_break_before(paragraph) -> None:
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph._p.insert(0, run)


def apply_paragraph_fonts(paragraph, *, is_heading: bool = False) -> None:
    if not has_cjk(paragraph.text):
        return
    east_asia = CJK_HEADING_FONT if is_heading else CJK_BODY_FONT
    for run in paragraph.runs:
        if run.text:
            set_east_asia_font(run, east_asia)


def write_docx(md_text: str, output: Path, title: str | None = None) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    set_style_font(style, CJK_BODY_FONT)
    style.font.size = Pt(11)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], CJK_HEADING_FONT)
    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], CJK_BODY_FONT)

    if title:
        p = doc.add_heading(title, level=0)
        apply_paragraph_fonts(p, is_heading=True)

    lines = md_text.splitlines()
    current_list = False
    wrote_content = bool(title)
    wrote_non_cjk_content = bool(title and not has_cjk(title))
    inserted_cjk_appendix_break = False

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            current_list = False
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            if has_cjk(stripped) and wrote_non_cjk_content and not inserted_cjk_appendix_break:
                add_page_break_before(p)
                inserted_cjk_appendix_break = True
            apply_paragraph_fonts(p, is_heading=True)
            current_list = False
            wrote_content = True
            wrote_non_cjk_content = wrote_non_cjk_content or not has_cjk(stripped)
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            apply_paragraph_fonts(p, is_heading=True)
            current_list = False
            wrote_content = True
            wrote_non_cjk_content = wrote_non_cjk_content or not has_cjk(stripped)
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            apply_paragraph_fonts(p, is_heading=True)
            current_list = False
            wrote_content = True
            wrote_non_cjk_content = wrote_non_cjk_content or not has_cjk(stripped)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            apply_paragraph_fonts(p)
            current_list = True
            wrote_content = True
            wrote_non_cjk_content = wrote_non_cjk_content or not has_cjk(stripped)
            continue
        if stripped[:2].isdigit() and stripped[1:3] == ". ":
            p = doc.add_paragraph(stripped[3:].strip(), style="List Number")
            apply_paragraph_fonts(p)
            current_list = True
            wrote_content = True
            wrote_non_cjk_content = wrote_non_cjk_content or not has_cjk(stripped)
            continue

        p = doc.add_paragraph()
        p.add_run(stripped)
        apply_paragraph_fonts(p)
        current_list = False
        wrote_content = True
        wrote_non_cjk_content = wrote_non_cjk_content or not has_cjk(stripped)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert markdown-like summary text into a Word document.")
    parser.add_argument("--input", required=True, type=Path, help="Input markdown/plain text file")
    parser.add_argument("--output", required=True, type=Path, help="Output .docx file")
    parser.add_argument("--title", help="Optional document title")
    args = parser.parse_args()

    text = args.input.read_text(encoding="utf-8")
    write_docx(text, args.output, args.title)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
