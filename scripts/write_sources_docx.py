#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert asset/source manifest JSON into a Word document.")
    parser.add_argument("--input", required=True, type=Path, help="Input JSON manifest")
    parser.add_argument("--output", required=True, type=Path, help="Output Word document")
    parser.add_argument("--title", default="来源清单", help="Document title")
    args = parser.parse_args()

    data = json.loads(args.input.read_text(encoding="utf-8"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    doc.add_heading(args.title, level=0)

    for idx, item in enumerate(data, start=1):
        doc.add_heading(f"{idx}. {item.get('title', '未命名来源')}", level=2)
        fields = [
            ("来源类型", item.get("source_type", "")),
            ("链接", item.get("url", "")),
            ("说明", item.get("license_or_note", "")),
            ("本地路径", item.get("local_path", "")),
            ("使用页码", str(item.get("used_for_slide", ""))),
        ]
        for label, value in fields:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            p.add_run(value)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
